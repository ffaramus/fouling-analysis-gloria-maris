import os
import cv2
import numpy as np
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans
import math
import matplotlib.cm as cm
from docx import Document
from docx.shared import Inches

# Variables globales pour stocker les points de la ROI
points = []

# Fonction de rappel pour capturer les clics de la souris
def onclick(event):
    global points
    if event.xdata is not None and event.ydata is not None:
        points.append((event.xdata, event.ydata))
        plt.plot(event.xdata, event.ydata, 'go')
        plt.draw()
        if len(points) == 4:
            plt.title("Appuyez sur 'entrer' pour continuer")
            plt.draw()

# Fonction pour corriger la distorsion de l'image
def correct_distortion(image):
    global points
    points = []

    def on_key(event):
        if event.key == 'enter':
            plt.close()

    fig, ax = plt.subplots()
    ax.imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
    fig.suptitle("Selectionnez 4 coins (Appuyez sur 'entrer' quand c'est fait)")
    fig.canvas.mpl_connect('button_press_event', onclick)
    fig.canvas.mpl_connect('key_press_event', on_key)
    plt.show()

    if len(points) != 4:
        print(f"Points sélectionnés: {len(points)}")
        raise ValueError("Quatre points doivent être sélectionnés.")

    source = np.float32(points)
    width = max(np.linalg.norm(source[0] - source[1]), np.linalg.norm(source[2] - source[3]))
    height = max(np.linalg.norm(source[0] - source[2]), np.linalg.norm(source[1] - source[3]))
    destination = np.float32([[0, 0], [width, 0], [0, height], [width, height]])
    matrix = cv2.getPerspectiveTransform(source, destination)
    corrected_image = cv2.warpPerspective(image, matrix, (int(width), int(height)))

    plt.imshow(cv2.cvtColor(corrected_image, cv2.COLOR_BGR2RGB))
    plt.title('Image selectionnée')
    plt.axis('off')
    plt.show()

    return corrected_image

# Fonction pour débrumer l'image
def dehaze_image(image):
    hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
    h, s, v = cv2.split(hsv)
    lim = 255 - 60
    v[v > lim] = 255
    v[v <= lim] += 60
    final_hsv = cv2.merge((h, s, v))
    dehazed_image = cv2.cvtColor(final_hsv, cv2.COLOR_HSV2BGR)

    plt.imshow(cv2.cvtColor(dehazed_image, cv2.COLOR_BGR2RGB))
    plt.title('Image débrumée')
    plt.axis('off')
    plt.savefig('image_débrumée.png')
    plt.close()
    return dehazed_image

# Fonction pour segmenter l'image
def segment_image(image, k=3):
    pixel_values = image.reshape((-1, 3))
    pixel_values = np.float32(pixel_values)
    kmeans = KMeans(n_clusters=k, init='k-means++', random_state=42)
    labels = kmeans.fit_predict(pixel_values)
    centers = kmeans.cluster_centers_
    centers = np.uint8(centers)
    segmented_image = centers[labels.flatten()]
    segmented_image = segmented_image.reshape(image.shape)

    plt.figure(figsize=(12, 6))
    plt.subplot(1, 2, 1)
    plt.imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
    plt.title('Image débrumée')
    plt.axis('off')

    plt.subplot(1, 2, 2)
    plt.imshow(cv2.cvtColor(segmented_image, cv2.COLOR_BGR2RGB))
    plt.title('Image segmentée')
    plt.axis('off')

    plt.tight_layout()
    plt.show()

    return segmented_image, labels.reshape(image.shape[:2])

# Fonction pour charger une image
def load_image(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Le fichier image à l'emplacement '{file_path}' n'a pas pu être chargé.")

    image = cv2.imread(file_path)
    if image is None:
        raise FileNotFoundError(f"Le fichier image à l'emplacement '{file_path}' n'a pas pu être chargé.")

    plt.imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
    plt.title('Image chargée')
    plt.axis('off')
    plt.show()

    return image

def overlap_area(rect1, rect2):
    x1, y1, w1, h1 = rect1
    x2, y2, w2, h2 = rect2

    # Coordonnées des coins des rectangles
    x_left = max(x1, x2)
    y_top = max(y1, y2)
    x_right = min(x1 + w1, x2 + w2)
    y_bottom = min(y1 + h1, y2 + h2)

    # Calcul de la largeur et de la hauteur de la zone de chevauchement
    overlap_width = max(0, x_right - x_left)
    overlap_height = max(0, y_bottom - y_top)

    # Surface de la zone de chevauchement
    return overlap_width * overlap_height


# Fonctions pour analyser la structure de l'ouverture du filet
def adjust_rectangles(rectangles, sr, MAX_X, MAX_Y):
    # 1. Filtrage des petits rectangles
    areas = [w * h for x, y, w, h in rectangles]
    mean_area = np.mean(areas)
    std_area = np.std(areas)

    # Conserver les rectangles qui sont suffisamment grands ou dont la taille est significative par rapport à la moyenne
    filtered_rectangles = [(x, y, w, h) for x, y, w, h in rectangles if w * h >= 0.4 * mean_area or w * h >= mean_area - std_area]

    if not filtered_rectangles:
        return []

    # 2. Calcul des centres des rectangles
    centers = [(x + w / 2, y + h / 2) for x, y, w, h in filtered_rectangles]
    idx_col = {i: [] for i in range(len(filtered_rectangles))}
    idx_row = {i: [] for i in range(len(filtered_rectangles))}

    for i, (cx1, cy1) in enumerate(centers):
        for j, (cx2, cy2) in enumerate(centers):
            if max(cx2 - sr, 0) <= cx1 <= min(cx2 + sr, MAX_X):
                idx_col[i].append(j)
            if max(cy2 - sr, 0) <= cy1 <= min(cy2 + sr, MAX_Y):
                idx_row[i].append(j)

    # 3. Ajuster les dimensions des rectangles pour qu'elles soient uniformes
    max_width = max(filtered_rectangles, key=lambda r: r[2])[2]
    max_height = max(filtered_rectangles, key=lambda r: r[3])[3]

    for i, (x, y, w, h) in enumerate(filtered_rectangles):
        filtered_rectangles[i] = (x, y, max_width, max_height)

    # 4. Réajuster les centres des rectangles
    optimized_rectangles = []
    for i in range(len(filtered_rectangles)):
        x, y, w, h = filtered_rectangles[i]
        if idx_col[i]:
            mean_cx = np.mean([centers[j][0] for j in idx_col[i]])
        else:
            mean_cx = centers[i][0]

        if idx_row[i]:
            mean_cy = np.mean([centers[j][1] for j in idx_row[i]])
        else:
            mean_cy = centers[i][1]

        optimized_rectangles.append((int(mean_cx - w / 2), int(mean_cy - h / 2), w, h))

    # 5. Ajouter les rectangles manquants
    col_centers = {}
    row_centers = {}

    for i in range(len(filtered_rectangles)):
        mean_cx = np.mean([centers[j][0] for j in idx_col[i]]) if idx_col[i] else centers[i][0]
        mean_cy = np.mean([centers[j][1] for j in idx_row[i]]) if idx_row[i] else centers[i][1]

        col_centers.setdefault(mean_cx, []).append(mean_cy)
        row_centers.setdefault(mean_cy, []).append(mean_cx)

    new_rectangles = set()
    for mean_cx in col_centers:
        for mean_cy in row_centers:
            # Vérification pour éviter le chevauchement excessif
            exists = any(abs(mean_cx - (x + w / 2)) < sr and abs(mean_cy - (y + h / 2)) < sr for x, y, w, h in optimized_rectangles)
            if not exists:
                Wmean = max_width
                Hmean = max_height
                new_rectangles.add((int(mean_cx - Wmean / 2), int(mean_cy - Hmean / 2), Wmean, Hmean))
                # Ne pas ajouter un rectangle s'il touche une bordure
                if (mean_cx - Wmean / 2) < 0 or (mean_cy - Hmean / 2) < 0 or (mean_cx + Wmean / 2) > MAX_X or (mean_cy + Hmean / 2) > MAX_Y:
                    continue
                new_rectangles.add((int(mean_cx - Wmean / 2), int(mean_cy - Hmean / 2), Wmean, Hmean))
    optimized_rectangles.extend(new_rectangles)

    # 6. Filtrer les rectangles qui se chevauchent excessivement
    final_rectangles = []
    for rect in optimized_rectangles:
        overlaps = [r for r in optimized_rectangles if overlap_area(rect, r) > 0.33 * rect[2] * rect[3]]
        if len(overlaps) <= 1:  # Conserver si peu de chevauchement
            final_rectangles.append(rect)

    return final_rectangles



def calculate_coverage(binary_image, rectangles):
    coverage_percentages = []

    for x, y, w, h in rectangles:
        rect_area = w * h
        roi = binary_image[y:y+h, x:x+w]
        black_area = np.sum(roi == 0)
        coverage_percentage = (black_area / rect_area) * 100
        coverage_percentages.append(coverage_percentage)

    mean_coverage = np.mean(coverage_percentages)
    return coverage_percentages, mean_coverage

def la_big_fonction(segmented_image, sr=30, MAX_X=2000, MAX_Y=2000):
    # # Redimensionner l'image pour une analyse plus fine
    # scale_factor = 2
    # resized_image = cv2.resize(segmented_image, None, fx=scale_factor, fy=scale_factor, interpolation=cv2.INTER_LINEAR)

    # Extraire les canaux de couleur et créer le masque de l'eau
    blue = segmented_image[:, :, 0].astype(np.float32)
    green = segmented_image[:, :, 1].astype(np.float32)
    red = segmented_image[:, :, 2].astype(np.float32)
    water_mask = (blue + green) - red
    water_mask = (water_mask - water_mask.min()) / (water_mask.max() - water_mask.min()) * 255
    water_mask = water_mask.astype(np.uint8)

    # Seuil pour créer l'image binaire
    _, binary_image = cv2.threshold(water_mask, 240, 255, cv2.THRESH_BINARY)

    # Appliquer l'érosion et la dilatation
    kernel = np.ones((5, 5), np.uint8)
    opening = cv2.morphologyEx(binary_image, cv2.MORPH_OPEN, kernel)

    # **Affichage de l'image avec les bords érodés
    plt.figure(figsize=(8, 8))
    plt.imshow(opening, cmap='gray')
    plt.title('Image avec bords érodés')
    plt.axis('off')
    plt.show()

    # Trouver les contours et créer des rectangles englobants
    contours, _ = cv2.findContours(opening, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    bounding_rectangles = []

    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        bounding_rectangles.append((x, y, w, h))

    # Afficher l'image avec les rectangles englobants
    rect_image = cv2.cvtColor(segmented_image, cv2.COLOR_BGR2RGB).copy()
    for (x, y, w, h) in bounding_rectangles:
        cv2.rectangle(rect_image, (x, y), (x + w, y + h), (255, 0, 0), 2)  # Dessine les rectangles en rouge
    plt.figure(figsize=(12, 8))
    plt.imshow(rect_image)
    plt.title('Image avec rectangles englobants')
    plt.axis('off')
    plt.show()

    # # Ajuster les rectangles
    # adjusted_rectangles = adjust_rectangles(bounding_rectangles, sr * scale_factor, MAX_X * scale_factor, MAX_Y * scale_factor)

    # Ajuster les rectangles
    adjusted_rectangles = adjust_rectangles(bounding_rectangles, sr, MAX_X, MAX_Y)

    # Afficher les rectangles ajustés
    adjusted_rect_image = cv2.cvtColor(segmented_image, cv2.COLOR_BGR2RGB).copy()
    for (x, y, w, h) in adjusted_rectangles:
        cv2.rectangle(adjusted_rect_image, (x, y), (x + w, y + h), (0, 255, 0), 2)  # Dessine les rectangles ajustés en vert
    plt.figure(figsize=(12, 8))
    plt.imshow(adjusted_rect_image)
    plt.title('Image avec rectangles ajustés')
    plt.axis('off')
    plt.show()

    # Calculer le pourcentage de recouvrement
    coverage_percentages, mean_coverage = calculate_coverage(binary_image, adjusted_rectangles)

    # Créer une image colorée en fonction des pourcentages de recouvrement
    colored_image = segmented_image.copy()
    cmap = cm.get_cmap('coolwarm_r')
    norm = plt.Normalize(0, 100)

    for i, (x, y, w, h) in enumerate(adjusted_rectangles):
        coverage = coverage_percentages[i]
        color = cmap(norm(coverage))[:3]
        color = tuple(int(c * 255) for c in color)
        cv2.rectangle(colored_image, (x, y), (x + w, y + h), color, cv2.FILLED)
        font_scale = max(0.5, min(w, h) / 100.0)
        cv2.putText(colored_image, f"{coverage:.2f}%", (x + w // 2 - 20, y + h // 2), cv2.FONT_HERSHEY_SIMPLEX, 0.5*font_scale, (255, 255, 255), 1, cv2.LINE_AA)

    # Afficher l'image avec les rectangles colorés et les pourcentages
    plt.figure(figsize=(12, 8))
    plt.imshow(cv2.cvtColor(colored_image, cv2.COLOR_BGR2RGB))
    plt.title(f'Analyse des ouvertures du filet avec rectangles ajustés\nRecouvrement moyen: {mean_coverage:.2f}%')
    plt.axis('off')
    plt.show()

    return colored_image, adjusted_rectangles, water_mask, binary_image

def create_word_report(corrected_image, water_mask_image, final_image, mean_coverage):
    # Créer un nouveau document Word
    doc = Document()

    # Ajouter le titre
    doc.add_heading('Analyse du filet 5d', level=1)

    # Ajouter les images avec les titres appropriés
    doc.add_heading('Portion sélectionnée', level=2)
    corrected_image_path = 'corrected_image.jpg'
    cv2.imwrite(corrected_image_path, corrected_image)
    doc.add_picture(corrected_image_path, width=Inches(4))

    doc.add_heading('Portion analysée', level=2)
    water_mask_image_path = 'water_mask_image.jpg'
    cv2.imwrite(water_mask_image_path, water_mask_image)
    doc.add_picture(water_mask_image_path, width=Inches(4))

    doc.add_heading('Résultat final', level=2)
    final_image_path = 'final_image.jpg'
    cv2.imwrite(final_image_path, final_image)
    doc.add_picture(final_image_path, width=Inches(4))

    # Ajouter le pourcentage moyen de recouvrement
    doc.add_paragraph(f'Recouvrement moyen: {mean_coverage:.2f}%')

    # Enregistrer le document
    doc_path = 'analyse_du_filet5d.docx'
    doc.save(doc_path)

    return doc_path

if __name__ == "__main__":
    # Charger une image
    image_path = 'C:/Users/MAR/Downloads/5/GOPR0805.jpg'
    image = cv2.imread(image_path)

    if image is None:
        print(f"Impossible de charger l'image à l'emplacement '{image_path}'.")
    else:
        corrected_image = correct_distortion(image)
        dehazed_image = dehaze_image(corrected_image)
        segmented_image, _ = segment_image(dehazed_image)
        net_structure_image, rectangles, water_mask, binary_image = la_big_fonction(segmented_image)

        if not rectangles:
            print("Aucun rectangle ajusté n'a été trouvé.")
        else:
            print(f"{len(rectangles)} rectangles ajustés ont été trouvés.")

            # Calculer le pourcentage moyen de recouvrement pour l'image finale
            _, mean_coverage = calculate_coverage(binary_image, rectangles)

            # Créer le rapport Word
            doc_path = create_word_report(corrected_image, water_mask, net_structure_image, mean_coverage)
            print(f"Le rapport a été créé et sauvegardé à: {doc_path}")


